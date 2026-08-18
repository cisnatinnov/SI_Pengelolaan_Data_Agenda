#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to generate UML_Documentation.docx from the UML documentation content.
The UML documentation is kept as a standalone document (UML_Documentation.docx)
and is NOT appended into skripsi.docx.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import struct

BASE_DIR = Path(__file__).resolve().parent
DIAGRAMS_DIR = BASE_DIR / "laravel-vue-mvc/docs/diagrams"
MAX_IMG_W = 6.0   # inches
MAX_IMG_H = 8.2   # inches


def add_heading_with_style(doc, text, level):
    """Add a heading with proper styling"""
    # Use level 1 for main title if level is 0
    if level == 0:
        level = 1
    heading = doc.add_heading(text, level=level)
    return heading


def add_table_from_markdown(doc, rows_data, header=None):
    """Add a table from markdown-style data"""
    num_cols = len(header) if header else len(rows_data[0])
    table = doc.add_table(rows=len(rows_data) + (1 if header else 0), cols=num_cols)
    table.style = 'Normal Table'

    row_idx = 0
    if header:
        header_cells = table.rows[0].cells
        for i, cell_text in enumerate(header):
            header_cells[i].text = cell_text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        row_idx = 1

    for row_data in rows_data:
        row_cells = table.rows[row_idx].cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
        row_idx += 1

    set_table_borders(table)
    return table


def add_uat_table(doc, title, rows):
    add_heading_with_style(doc, title, 2)
    add_table_from_markdown(doc, rows, header=["ID", "Skenario Uji", "Langkah / Input", "Hasil yang Diharapkan", "Hasil Aktual", "Status"])


def set_table_borders(table, color="808080", sz=6):
    """Apply a single-line border to every cell of the table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)
    table.autofit = True


def png_size(path):
    """Return (width, height) in pixels of a PNG file."""
    with open(path, "rb") as fh:
        head = fh.read(26)
    if head[:8] != b"\x89PNG\r\n\x1a\n":
        raise ValueError(f"Not a PNG: {path}")
    w, h = struct.unpack(">II", head[16:24])
    return w, h


def add_diagram(doc, filename, caption):
    """Insert a rendered Mermaid diagram (PNG) centered with a caption."""
    img_path = DIAGRAMS_DIR / filename
    w, h = png_size(img_path)

    scale = min(MAX_IMG_W / w, MAX_IMG_H / h, 1.0)
    width = w * scale
    height = h * scale

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width), height=Inches(height))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def main():
    # Create a fresh standalone document.
    doc = Document()

    # Add Title
    add_heading_with_style(doc, "Dokumentasi UML — SI Pengelolaan Data Agenda", 0)

    # Add introduction
    doc.add_paragraph(
        "Dokumen ini berisi diagram UML (Use Case, Sequence, Activity, Robustness, dan Class) untuk sistem "
        "SI Pengelolaan Data Agenda, aplikasi berbasis Laravel 12 + Vue 3 (PWA) untuk mengelola "
        "surat, disposisi, kegiatan, kehadiran, dan pengingat."
    )

    # Ringkasan Sistem
    add_heading_with_style(doc, "Ringkasan Sistem", 1)

    sistem_items = [
        "Login menggunakan email + password (session-based).",
        "Role: Admin, Staff, Asisten Daerah, OPD.",
        "Saat Staff membuat Surat, sistem otomatis:\n   1. Membuat Disposisi dengan keterangan diterima (tanpa perlu diedit), dan\n   2. Membuat Pengingat untuk seluruh akun Asisten Daerah.",
        "Saat salah satu role menambah Kegiatan, sistem otomatis membuat Pengingat untuk semua role.",
        "Saat menambah Kegiatan, sistem otomatis melakukan Cek Bentrok Jadwal: jika sudah ada kegiatan lain pada tanggal dan jam yang sama, pembuatan kegiatan ditolak.",
        "Hanya OPD yang dapat melakukan Konfirmasi Kehadiran kegiatan (hadir / tidak), tercatat per akun OPD. Role lain dapat melihat rekap dan daftar OPD yang mengonfirmasi.",
        "Disposisi dapat Diserahkan atau Ditolak (wajib alasan) oleh Asisten Daerah. Saat Diserahkan, Asisten Daerah wajib mengisi Penerima (tandatangan_penerima) dan Dituju (tandatangan_dituju). Sistem otomatis membuat Pengingat untuk seluruh akun Staff.",
        "Staff hanya dapat melihat Disposisi: tidak dapat mengubah, menyerahkan/menolak, maupun menghapus. Hanya Asisten Daerah yang dapat menyerahkan/menolak; tidak ada role yang dapat menghapus disposisi.",
        "Lonceng Notifikasi Pengingat tersedia di header untuk role Staff, Asisten Daerah, dan OPD.",
        "Notifikasi Real-Time (via Laravel Reverb + Laravel Echo) hanya aktif untuk pengingat yang dibuat otomatis dari form Tambah Surat (source = surat), Tambah Kegiatan (source = kegiatan), dan aksi Menyerahkan/Menolak Disposisi (source = disposisi). Pengingat yang dibuat manual tidak memicu notifikasi.",
        "Setiap pengingat otomatis dilengkapi status dibaca / belum dibaca (read_at) dan dapat ditandai dibaca per item atau semuanya."
    ]

    for item in sistem_items:
        p = doc.add_paragraph(item, style='List Paragraph')
        p.paragraph_format.left_indent = Inches(0.25)

    # Aktor table
    add_heading_with_style(doc, "Aktor", 1)
    aktor_data = [
        ["Admin", "Mengelola pengguna & role. Tidak memiliki akses pengingat/notifikasi."],
        ["Staff", "Mengelola surat dan kegiatan, melihat disposisi, dan mengelola pengingat pribadi; menerima notifikasi pengingat real-time."],
        ["Asisten Daerah", "Meninjau disposisi (menyerahkan/menolak), mengelola pengingat pribadi, dan menerima notifikasi pengingat real-time."],
        ["OPD", "Mengonfirmasi kehadiran kegiatan, mengelola pengingat pribadi, dan menerima notifikasi pengingat real-time."]
    ]
    add_table_from_markdown(doc, aktor_data, header=["Aktor", "Deskripsi"])

    # Kebutuhan Fungsional Sistem
    add_heading_with_style(doc, "Kebutuhan Fungsional Sistem", 1)
    fr_data = [
        ["FR-01", "Login & Logout", "Autentikasi pengguna menggunakan email + password (session-based) beserta logout.", "Admin, Staff, Asisten Daerah, OPD"],
        ["FR-02", "Dashboard", "Menampilkan statistik Disposisi (Total, Diterima, Ditolak, Diserahkan), statistik Kegiatan (Total, Terlaksana, Tidak), dan rekap Kegiatan per Periode.", "Admin, Staff, Asisten Daerah, OPD"],
        ["FR-03", "Kelola Surat", "Membuat, mengubah, dan menghapus surat. Saat surat baru dibuat, sistem otomatis membuat Disposisi (status Diterima) dan Pengingat untuk seluruh Asisten Daerah.", "Staff"],
        ["FR-04", "Lihat Disposisi", "Melihat daftar dan detail disposisi beserta status dan alasan.", "Staff, Asisten Daerah"],
        ["FR-05", "Serahkan Disposisi", "Menyerahkan disposisi disertai Penerima (tandatangan_penerima) dan Dituju (tandatangan_dituju) tanpa alasan. Sistem otomatis membuat Pengingat ke seluruh Staff.", "Asisten Daerah"],
        ["FR-06", "Tolak Disposisi", "Menolak disposisi disertai Alasan (wajib). Sistem otomatis membuat Pengingat ke seluruh Staff.", "Asisten Daerah"],
        ["FR-07", "Kelola Kegiatan", "Membuat, mengubah, dan menghapus kegiatan dengan Cek Bentrok Jadwal otomatis (menolak jika tanggal & jam sama). Sistem otomatis membuat Pengingat ke semua role.", "Staff"],
        ["FR-08", "Konfirmasi Kehadiran", "Mengonfirmasi kehadiran kegiatan (Hadir / Tidak) yang tercatat per akun OPD dan dapat diubah.", "OPD"],
        ["FR-09", "Lihat Rekap & Daftar OPD", "Melihat rekap jumlah hadir/tidak dan daftar OPD yang mengonfirmasi kehadiran suatu kegiatan.", "Staff, Asisten Daerah"],
        ["FR-10", "Kelola Pengingat Pribadi", "Membuat, mengubah, menghapus, dan menandai selesai pengingat milik sendiri.", "Staff, Asisten Daerah, OPD"],
        ["FR-11", "Notifikasi Pengingat Real-Time", "Menerima notifikasi real-time (badge belum dibaca) saat surat/kegiatan/disposisi dibuat atau diubah, serta menandai dibaca per item atau semuanya.", "Staff, Asisten Daerah, OPD"],
        ["FR-12", "Kelola Pengguna & Role", "Mengelola akun pengguna dan role.", "Admin"],
    ]
    add_table_from_markdown(doc, fr_data, header=["Kode", "Kebutuhan Fungsional", "Deskripsi", "Aktor"])

    # Use Case Diagram
    add_heading_with_style(doc, "1. Use Case Diagram", 1)
    doc.add_paragraph(
        "Catatan: <<include>> menandai proses tambahan yang dijalankan otomatis oleh sistem setelah aksi utama."
    )
    add_diagram(doc, "01_1_use_case_diagram.png", "Gambar: Use Case Diagram")

    # 1.1 Use Case Login & Dashboard
    add_heading_with_style(doc, "1.1 Use Case — Login dan Dashboard", 2)
    doc.add_paragraph(
        "Use case ini memfokuskan pada dua fungsionalitas yang diakses oleh semua role (Admin, Staff, "
        "Asisten Daerah, OPD), yaitu Login/Logout dan Dashboard."
    )
    add_diagram(doc, "02_1_1_use_case_login_dan_dashboard.png", "Gambar: Use Case Login dan Dashboard")

    # Sequence Diagrams
    add_heading_with_style(doc, "2. Sequence Diagram", 1)

    # 2.1 Login & Logout (FR-01)
    add_heading_with_style(doc, "2.1 Login & Logout (Autentikasi Pengguna) — FR-01", 2)
    add_diagram(doc, "03_2_1_login_logout_autentikasi_pengguna_fr_01.png", "Gambar: Sequence Diagram Login & Logout")

    # 2.2 Dashboard (FR-02)
    add_heading_with_style(doc, "2.2 Dashboard (Melihat Statistik Disposisi & Kegiatan) — FR-02", 2)
    add_diagram(doc, "04_2_2_dashboard_melihat_statistik_disposisi_kegiatan_fr_02.png", "Gambar: Sequence Diagram Dashboard")

    # 2.3 Kelola Surat (FR-03)
    add_heading_with_style(doc, "2.3 Kelola Surat (Buat / Ubah / Hapus) — Auto Disposisi & Pengingat — FR-03", 2)
    add_diagram(doc, "05_2_3_kelola_surat_buat_ubah_hapus_auto_disposisi_pengingat_fr_03.png", "Gambar: Sequence Diagram Kelola Surat")

    # 2.4 Lihat Disposisi (FR-04)
    add_heading_with_style(doc, "2.4 Lihat Disposisi (Daftar & Detail) — FR-04", 2)
    add_diagram(doc, "06_2_4_lihat_disposisi_fr_04.png", "Gambar: Sequence Diagram Lihat Disposisi")

    # 2.5 Asisten Meninjau Disposisi (FR-05 & FR-06)
    add_heading_with_style(doc, "2.5 Asisten Daerah Meninjau Disposisi (Serahkan / Tolak) — FR-05 & FR-06", 2)
    add_diagram(doc, "07_2_5_asisten_daerah_meninjau_disposisi_serahkan_tolak_pengingat_staff_fr_05_fr_06.png", "Gambar: Sequence Diagram Asisten Daerah Meninjau Disposisi")

    # 2.6 Kelola Kegiatan (FR-07)
    add_heading_with_style(doc, "2.6 Kelola Kegiatan (Buat / Ubah / Hapus) dengan Auto Cek Bentrok Jadwal — FR-07", 2)
    add_diagram(doc, "08_2_6_kelola_kegiatan_buat_ubah_hapus_dengan_auto_cek_bentrok_jadwal_fr_07.png", "Gambar: Sequence Diagram Kelola Kegiatan")

    # 2.7 OPD Konfirmasi Kehadiran (FR-08)
    add_heading_with_style(doc, "2.7 OPD Konfirmasi Kehadiran Kegiatan — FR-08", 2)
    add_diagram(doc, "09_2_7_opd_konfirmasi_kehadiran_kegiatan_fr_08.png", "Gambar: Sequence Diagram Konfirmasi Kehadiran")

    # 2.8 Lihat Rekap & Daftar OPD (FR-09)
    add_heading_with_style(doc, "2.8 Lihat Rekap & Daftar OPD — FR-09", 2)
    add_diagram(doc, "10_2_8_lihat_rekap_daftar_opd_fr_09.png", "Gambar: Sequence Diagram Lihat Rekap & Daftar OPD")

    # 2.9 Kelola Pengingat Pribadi (FR-10)
    add_heading_with_style(doc, "2.9 Kelola Pengingat Pribadi — FR-10", 2)
    add_diagram(doc, "11_2_9_kelola_pengingat_pribadi_fr_10.png", "Gambar: Sequence Diagram Kelola Pengingat Pribadi")

    # 2.10 Notifikasi Pengingat Real-Time (FR-11)
    add_heading_with_style(doc, "2.10 Notifikasi Pengingat Real-Time (Lonceng Notifikasi) — FR-11", 2)
    add_diagram(doc, "12_2_10_notifikasi_pengingat_real_time_lonceng_notifikasi_fr_11.png", "Gambar: Sequence Diagram Notifikasi Pengingat Real-Time")

    # 2.11 Kelola Pengguna & Role (FR-12)
    add_heading_with_style(doc, "2.11 Kelola Pengguna & Role — FR-12", 2)
    add_diagram(doc, "13_2_11_kelola_pengguna_role_fr_12.png", "Gambar: Sequence Diagram Kelola Pengguna & Role")

    # Activity Diagrams
    add_heading_with_style(doc, "3. Activity Diagram", 1)
    doc.add_paragraph(
        "Setiap use case pada Use Case Diagram dipetakan ke sebuah Activity Diagram (satu activity diagram "
        "per use case / kebutuhan fungsional)."
    )

    # 3.1 Login & Logout (FR-01)
    add_heading_with_style(doc, "3.1 Alur Login & Logout — FR-01", 2)
    add_diagram(doc, "14_3_1_alur_login_logout_fr_01.png", "Gambar: Activity Diagram Login & Logout")

    # 3.2 Dashboard (FR-02)
    add_heading_with_style(doc, "3.2 Alur Menampilkan Dashboard — FR-02", 2)
    add_diagram(doc, "15_3_2_alur_menampilkan_dashboard_fr_02.png", "Gambar: Activity Diagram Dashboard")

    # 3.3 Kelola Surat (FR-03)
    add_heading_with_style(doc, "3.3 Alur Kelola Surat — FR-03", 2)
    add_diagram(doc, "16_3_3_alur_kelola_surat_fr_03.png", "Gambar: Activity Diagram Kelola Surat")

    # 3.4 Lihat Disposisi (FR-04)
    add_heading_with_style(doc, "3.4 Alur Lihat Disposisi — FR-04", 2)
    add_diagram(doc, "17_3_4_alur_lihat_disposisi_fr_04.png", "Gambar: Activity Diagram Lihat Disposisi")

    # 3.5 Menyerahkan Disposisi (FR-05)
    add_heading_with_style(doc, "3.5 Alur Menyerahkan Disposisi — FR-05", 2)
    add_diagram(doc, "18_3_5_alur_menyerahkan_disposisi_fr_05.png", "Gambar: Activity Diagram Menyerahkan Disposisi")

    # 3.6 Menolak Disposisi (FR-06)
    add_heading_with_style(doc, "3.6 Alur Menolak Disposisi — FR-06", 2)
    add_diagram(doc, "19_3_6_alur_menolak_disposisi_fr_06.png", "Gambar: Activity Diagram Menolak Disposisi")

    # 3.7 Kelola Kegiatan (FR-07)
    add_heading_with_style(doc, "3.7 Alur Kelola Kegiatan — FR-07", 2)
    add_diagram(doc, "20_3_7_alur_kelola_kegiatan_fr_07.png", "Gambar: Activity Diagram Kelola Kegiatan")

    # 3.8 Konfirmasi Kehadiran (FR-08)
    add_heading_with_style(doc, "3.8 Alur Konfirmasi Kehadiran Kegiatan (OPD) — FR-08", 2)
    add_diagram(doc, "21_3_8_alur_konfirmasi_kehadiran_kegiatan_opd_fr_08.png", "Gambar: Activity Diagram Konfirmasi Kehadiran")

    # 3.9 Lihat Rekap & Daftar OPD (FR-09)
    add_heading_with_style(doc, "3.9 Alur Lihat Rekap & Daftar OPD — FR-09", 2)
    add_diagram(doc, "22_3_9_alur_lihat_rekap_daftar_opd_fr_09.png", "Gambar: Activity Diagram Lihat Rekap & Daftar OPD")

    # 3.10 Kelola Pengingat Pribadi (FR-10)
    add_heading_with_style(doc, "3.10 Alur Kelola Pengingat Pribadi — FR-10", 2)
    add_diagram(doc, "23_3_10_alur_kelola_pengingat_pribadi_fr_10.png", "Gambar: Activity Diagram Kelola Pengingat Pribadi")

    # 3.11 Notifikasi Pengingat Real-Time (FR-11)
    add_heading_with_style(doc, "3.11 Alur Notifikasi Pengingat Real-Time — FR-11", 2)
    add_diagram(doc, "24_3_11_alur_notifikasi_pengingat_real_time_fr_11.png", "Gambar: Activity Diagram Notifikasi Pengingat Real-Time")

    # 3.12 Kelola Pengguna & Role (FR-12)
    add_heading_with_style(doc, "3.12 Alur Kelola Pengguna & Role — FR-12", 2)
    add_diagram(doc, "25_3_12_alur_kelola_pengguna_role_fr_12.png", "Gambar: Activity Diagram Kelola Pengguna & Role")

    # Robustness Diagram
    add_heading_with_style(doc, "4. Robustness Diagram", 1)
    doc.add_paragraph(
        "Robustness diagram memodelkan hubungan antar tiga jenis objek yang saling berinteraksi dalam setiap "
        "alur: Boundary (antarmuka pengguna), Controller (logika aplikasi), dan Entity (data/model). Diagram "
        "ini digunakan untuk memastikan setiap use case dapat dipetakan ke antarmuka, kontrol, dan data yang sesuai."
    )
    add_diagram(doc, "26_4_robustness_diagram.png", "Gambar: Robustness Diagram")

    # Class Diagram
    add_heading_with_style(doc, "5. Class Diagram", 1)
    add_diagram(doc, "27_5_class_diagram.png", "Gambar: Class Diagram")

    # Add classes description
    add_heading_with_style(doc, "Class Descriptions", 2)

    classes = {
        "Role": ["id (int)", "name (string)", "slug (string)"],
        "User": ["id (int)", "name (string)", "email (string)", "email_verified_at (datetime, nullable)", "role_id (int)", "password (string)", "remember_token (string, nullable)", "role_slug() (method)"],
        "Surat": ["id (int)", "tanggal (datetime)", "nomor_surat (string)", "asal_surat (string)", "perihal (string)",
                  "kepada (string)", "tanggal_pelaksanaan (datetime)", "tempat_pelaksanaan (string)",
                  "pembawa_surat (string)", "tandatangan (string)", "disposisis() (method)"],
        "Disposisi": ["id (int)", "surat_id (int)", "tanggal (datetime)", "nomor_surat (string)", "asal_surat (string)",
                      "perihal (string)", "kepada (string)", "pembawa_surat (string)", "tandatangan_penerima (string)",
                      "tandatangan_dituju (string)", "keterangan (string)", "alasan (string)", "surat() (method)"],
        "Kegiatan": ["id (int)", "nama_kegiatan (string)", "tempat_kegiatan (string)", "tanggal_kegiatan (datetime)",
                     "uraian_kegiatan (string)", "realisasi_pelaksanaan (string)", "keterangan (string)",
                     "status (string)", "nama_penyusun (string)", "kehadiran() (method)"],
        "KegiatanKehadiran": ["id (int)", "kegiatan_id (int)", "user_id (int)", "status (string)",
                              "kegiatan() (method)", "user() (method)"],
        "Pengingat": ["id (int)", "user_id (int)", "judul (string)", "deskripsi (string)",
                      "tanggal_pengingat (datetime)", "prioritas (string)", "status (string)",
                      "source (string: manual|surat|kegiatan|disposisi)", "read_at (datetime, nullable)",
                      "user() (method)"],
        "PengingatNotification": ["pengingat (Pengingat)", "broadcastOn() (method)", "broadcastAs() (method)",
                      "Catatan: event broadcast Laravel Reverb, bukan tabel database"]
    }

    for class_name, attributes in classes.items():
        p = doc.add_paragraph(style='List Paragraph')
        run = p.add_run(f"{class_name}: ")
        run.bold = True
        p.add_run(", ".join(attributes))
        p.paragraph_format.left_indent = Inches(0.25)

    # Relationships
    add_heading_with_style(doc, "Relationships", 2)
    relationships = [
        "Role (1) → (0..*) User : role_id",
        "User (1) → (0..*) Pengingat : user_id",
        "Surat (1) → (0..*) Disposisi : surat_id",
        "Kegiatan (1) → (0..*) KegiatanKehadiran : kegiatan_id",
        "User (1) → (0..*) KegiatanKehadiran : user_id",
        "PengingatNotification (1) → (1) Pengingat : pengingat"
    ]
    for rel in relationships:
        p = doc.add_paragraph(rel, style='List Paragraph')
        p.paragraph_format.left_indent = Inches(0.25)

    # Enum / Status values
    add_heading_with_style(doc, "Nilai Enum / Status", 1)

    enum_data = [
        ["Disposisi.keterangan", "diterima, ditolak, diserahkan"],
        ["Kegiatan.realisasi_pelaksanaan", "terlaksana, tidak"],
        ["Kegiatan.status", "pelaksanaan, laporan"],
        ["KegiatanKehadiran.status", "hadir, tidak"],
        ["Pengingat.prioritas", "rendah, sedang, tinggi"],
        ["Pengingat.status", "pending, selesai"],
        ["Pengingat.source", "manual, surat, kegiatan, disposisi"],
        ["Pengingat.read_at", "null (belum dibaca), timestamp (dibaca)"]
    ]
    add_table_from_markdown(doc, enum_data, header=["Atribut", "Nilai"])

    # 6. Rancangan Database
    add_heading_with_style(doc, "6. Rancangan Database", 1)
    doc.add_paragraph(
        "Rancangan database berikut menggambarkan struktur tabel fisik pada database SQLite aplikasi, "
        "disusun dari skema migrasi (Laravel Migration). Relasi antar tabel mengikuti kunci asing "
        "(foreign key) yang didefinisikan pada migrasi."
    )
    add_diagram(doc, "28_6_rancangan_database.png", "Gambar: Rancangan Database (ERD)")
    doc.add_paragraph(
        "Catatan: tabel cache, jobs, password_reset_tokens, dan sessions merupakan tabel bawaan framework "
        "Laravel dan tidak termasuk domain bisnis sistem, sehingga tidak digambarkan pada Rancangan Database."
    )

    # 7. Skenario Pengujian API — Response Berhasil / Gagal
    add_heading_with_style(doc, "7. Skenario Pengujian API — Response Berhasil / Gagal", 1)
    doc.add_paragraph(
        "Tabel berikut memetakan setiap use case ke skenario pengujian beserta kode response HTTP yang "
        "dikembalikan sistem, baik pada kondisi berhasil maupun gagal. Endpoint API mengembalikan format JSON; "
        "endpoint web (login/logout) mengembalikan redirect."
    )
    doc.add_paragraph(
        "Keterangan status: 200 sukses dengan data, 201 sukses buat data baru, 204 sukses hapus tanpa konten, "
        "302 redirect, 401 belum login, 403 role tidak berhak, 404 data tidak ditemukan, 422 validasi gagal."
    )

    api_tables = [
        ("7.1 Login & Dashboard", [
            ["A-01", "Membuka halaman login", "Guest", "GET /login", "200", "Berhasil — form login ditampilkan"],
            ["A-02", "Login", "Email & password benar", "POST /login", "302", "Berhasil — redirect ke /"],
            ["A-03", "Login", "Email/password salah", "POST /login", "302 + error", "Gagal — pesan 'Email atau password salah'"],
            ["A-04", "Login", "Input tidak valid / password lemah", "POST /login", "302 + error validasi", "Gagal — error validasi email & kekuatan password"],
            ["A-05", "Login (ingat saya)", "Ceklis 'remember me'", "POST /login", "302 + cookie", "Berhasil — sesi diingat"],
            ["A-06", "Logout", "User terautentikasi", "POST /logout", "302", "Berhasil — redirect ke /login"],
            ["A-07", "Logout", "Guest", "POST /logout", "302", "Gagal — redirect ke /login"],
            ["A-08", "Akses dashboard /", "Guest", "GET /", "302", "Gagal — redirect ke /login"],
            ["A-09", "Akses dashboard /", "User terautentikasi", "GET /", "200", "Berhasil — dashboard tampil"],
        ]),
        ("7.2 Kelola Surat", [
            ["S-01", "Lihat daftar surat", "Terautentikasi", "GET /api/surat", "200", "Berhasil — array surat"],
            ["S-02", "Lihat daftar surat", "Guest", "GET /api/surat", "401", "Gagal — belum login"],
            ["S-03", "Buat surat", "Data valid", "POST /api/surat", "201", "Berhasil — auto buat Disposisi (diterima) & Pengingat"],
            ["S-04", "Buat surat", "Field wajib kosong", "POST /api/surat", "422", "Gagal — error validasi"],
            ["S-05", "Detail surat", "Data tersedia", "GET /api/surat/{id}", "200", "Berhasil — data surat"],
            ["S-06", "Detail surat", "Data tidak ditemukan", "GET /api/surat/999", "404", "Gagal — surat tidak ada"],
            ["S-07", "Ubah surat", "Data valid", "PUT /api/surat/{id}", "200", "Berhasil — data terupdate"],
            ["S-08", "Ubah surat", "Field tidak valid", "PUT /api/surat/{id}", "422", "Gagal — error validasi"],
            ["S-09", "Hapus surat", "Data tersedia", "DELETE /api/surat/{id}", "204", "Berhasil — data terhapus"],
        ]),
        ("7.3 Kelola Disposisi (Lihat / Serahkan / Tolak)", [
            ["D-01", "Lihat daftar disposisi", "Terautentikasi", "GET /api/disposisi", "200", "Berhasil — array disposisi"],
            ["D-02", "Lihat disposisi per surat", "Filter ?surat_id=", "GET /api/disposisi?surat_id=5", "200", "Berhasil — hasil terfilter"],
            ["D-03", "Detail disposisi", "Data tersedia", "GET /api/disposisi/{id}", "200", "Berhasil — data + relasi surat"],
            ["D-04", "Detail disposisi", "Data tidak ditemukan", "GET /api/disposisi/999", "404", "Gagal — disposisi tidak ada"],
            ["D-05", "Staff ubah disposisi", "Bukan role asisten", "PUT /api/disposisi/{id}", "403", "Gagal — staff tidak memiliki akses mengubah"],
            ["D-06", "Staff serahkan/tolak", "Bukan role asisten", "PUT /api/disposisi/{id}", "403", "Gagal — staff tidak memiliki akses"],
            ["D-07", "Staff/Asisten hapus", "Route tidak terdaftar", "DELETE /api/disposisi/{id}", "405", "Gagal — method tidak diizinkan"],
            ["D-08", "Asisten menyerahkan", "keterangan = diserahkan + tandatangan_penerima + tandatangan_dituju", "PUT /api/disposisi/{id}", "200", "Berhasil — status Diserahkan + Penerima/Dituju tersimpan + Pengingat otomatis ke seluruh Staff"],
            ["D-09", "Asisten menolak", "keterangan = ditolak + alasan", "PUT /api/disposisi/{id}", "200", "Berhasil — status Ditolak + alasan + Pengingat otomatis ke seluruh Staff"],
            ["D-10", "Asisten menolak", "Tanpa alasan", "PUT /api/disposisi/{id}", "422", "Gagal — alasan wajib diisi saat ditolak"],
            ["D-10a", "Asisten menyerahkan", "Tanpa tandatangan_penerima/dituju", "PUT /api/disposisi/{id}", "422", "Gagal — Penerima & Dituju wajib diisi saat diserahkan"],
            ["D-11", "Asisten memakai field staff", "keterangan = diterima", "PUT /api/disposisi/{id}", "422", "Gagal — tidak diizinkan untuk role asisten"],
            ["D-12", "Role lain (OPD) ubah", "Bukan asisten", "PUT /api/disposisi/{id}", "403", "Gagal — tidak memiliki akses"],
        ]),
        ("7.4 Kelola Kegiatan (Buat / Ubah / Hapus) + Auto Cek Bentrok Jadwal", [
            ["K-01", "Lihat daftar kegiatan", "Semua role", "GET /api/kegiatan", "200", "Berhasil — daftar + hadir_count/tidak_count"],
            ["K-02", "Detail kegiatan", "Data tersedia", "GET /api/kegiatan/{id}", "200", "Berhasil — data kegiatan"],
            ["K-03", "Detail kegiatan", "Data tidak ditemukan", "GET /api/kegiatan/999", "404", "Gagal — kegiatan tidak ada"],
            ["K-04", "Buat kegiatan", "Staff, jadwal tersedia", "POST /api/kegiatan", "201", "Berhasil — auto Pengingat semua role"],
            ["K-05", "Buat kegiatan", "Jadwal bentrok", "POST /api/kegiatan", "422", "Gagal — jadwal bentrok"],
            ["K-06", "Buat kegiatan", "Field wajib kosong", "POST /api/kegiatan", "422", "Gagal — error validasi"],
            ["K-07", "Buat kegiatan", "Role bukan staff", "POST /api/kegiatan", "403", "Gagal — tidak memiliki akses"],
            ["K-08", "Ubah kegiatan", "Staff, jadwal tersedia", "PUT /api/kegiatan/{id}", "200", "Berhasil — data terupdate"],
            ["K-09", "Ubah kegiatan", "Jadwal bentrok", "PUT /api/kegiatan/{id}", "422", "Gagal — jadwal bentrok"],
            ["K-10", "Ubah kegiatan", "Mempertahankan jadwal sendiri", "PUT /api/kegiatan/{id}", "200", "Berhasil — tidak dianggap bentrok"],
            ["K-11", "Ubah kegiatan", "Role bukan staff", "PUT /api/kegiatan/{id}", "403", "Gagal — tidak memiliki akses"],
            ["K-12", "Hapus kegiatan", "Staff", "DELETE /api/kegiatan/{id}", "204", "Berhasil — data terhapus"],
            ["K-13", "Hapus kegiatan", "Role bukan staff", "DELETE /api/kegiatan/{id}", "403", "Gagal — tidak memiliki akses"],
        ]),
        ("7.5 Konfirmasi Kehadiran (OPD)", [
            ["H-01", "Konfirmasi hadir", "Role OPD, status = hadir", "POST /api/kegiatan/{id}/kehadiran", "200", "Berhasil — tercatat per akun OPD"],
            ["H-02", "Konfirmasi tidak hadir", "Role OPD, status = tidak", "POST /api/kegiatan/{id}/kehadiran", "200", "Berhasil — status diperbarui"],
            ["H-03", "Ubah konfirmasi", "OPD konfirmasi ulang", "POST /api/kegiatan/{id}/kehadiran", "200", "Berhasil — record di-update"],
            ["H-04", "Konfirmasi", "status tidak valid", "POST /api/kegiatan/{id}/kehadiran", "422", "Gagal — error validasi"],
            ["H-05", "Konfirmasi", "Role bukan OPD", "POST /api/kegiatan/{id}/kehadiran", "403", "Gagal — tidak memiliki akses"],
            ["H-06", "Konfirmasi", "Kegiatan tidak ditemukan", "POST /api/kegiatan/999/kehadiran", "404", "Gagal — kegiatan tidak ada"],
        ]),
        ("7.6 Kelola Pengingat Pribadi (Staff / Asisten Daerah / OPD)", [
            ["P-01", "Lihat daftar pengingat", "Data milik sendiri", "GET /api/pengingat", "200", "Berhasil — hanya milik sendiri"],
            ["P-02", "Buat pengingat", "Data valid", "POST /api/pengingat", "201", "Berhasil — tersimpan"],
            ["P-03", "Buat pengingat", "Field wajib kosong", "POST /api/pengingat", "422", "Gagal — error validasi"],
            ["P-04", "Detail pengingat", "Milik user yang sama", "GET /api/pengingat/{id}", "200", "Berhasil — data pengingat"],
            ["P-05", "Detail pengingat", "Milik user lain", "GET /api/pengingat/{id}", "404", "Gagal — tidak ditemukan"],
            ["P-06", "Ubah pengingat", "Milik user yang sama", "PUT /api/pengingat/{id}", "200", "Berhasil — data terupdate"],
            ["P-07", "Ubah pengingat", "Milik user lain", "PUT /api/pengingat/{id}", "404", "Gagal — tidak ditemukan"],
            ["P-08", "Ubah pengingat", "Data tidak valid", "PUT /api/pengingat/{id}", "422", "Gagal — error validasi"],
            ["P-09", "Hapus pengingat", "Milik user yang sama", "DELETE /api/pengingat/{id}", "204", "Berhasil — data terhapus"],
            ["P-10", "Hapus pengingat", "Milik user lain", "DELETE /api/pengingat/{id}", "404", "Gagal — tidak ditemukan"],
            ["P-11", "Akses pengingat", "Role Admin", "GET /api/pengingat", "403", "Gagal — admin tidak diperbolehkan"],
        ]),
        ("7.8 Notifikasi Pengingat Real-Time", [
            ["N-01", "Lihat notifikasi", "Terautentikasi", "GET /api/pengingat/notifications", "200", "Berhasil — hanya source = surat/kegiatan + unread_count"],
            ["N-02", "Tandai notifikasi dibaca", "Milik user yang sama", "POST /api/pengingat/{id}/read", "200", "Berhasil — read_at terisi"],
            ["N-03", "Tandai notifikasi dibaca", "Milik user lain", "POST /api/pengingat/{id}/read", "404", "Gagal — diperlakukan sebagai tidak ditemukan"],
            ["N-04", "Tandai semua notifikasi dibaca", "Terautentikasi", "POST /api/pengingat/read-all", "200", "Berhasil — semua read_at terisi"],
            ["N-05", "Akses notifikasi", "Role Admin", "GET /api/pengingat/notifications", "403", "Gagal — admin tidak diperbolehkan"],
        ]),
        ("7.7 Kelola Pengguna & Role (Admin)", [
            ["U-01", "Lihat daftar user", "Role Admin", "GET /api/users", "200", "Berhasil — daftar user + role"],
            ["U-02", "Buat user", "Data valid (password kuat)", "POST /api/users", "201", "Berhasil — user tersimpan"],
            ["U-03", "Buat user", "Email sudah terpakai", "POST /api/users", "422", "Gagal — email duplikat"],
            ["U-04", "Buat user", "Password lemah", "POST /api/users", "422", "Gagal — password tidak memenuhi syarat"],
            ["U-05", "Detail user", "Data tersedia", "GET /api/users/{id}", "200", "Berhasil — data user"],
            ["U-06", "Detail user", "Data tidak ditemukan", "GET /api/users/999", "404", "Gagal — user tidak ada"],
            ["U-07", "Ubah user", "Data valid", "PUT /api/users/{id}", "200", "Berhasil — data terupdate"],
            ["U-08", "Ubah user", "Data tidak valid", "PUT /api/users/{id}", "422", "Gagal — error validasi"],
            ["U-09", "Hapus user", "Menghapus user lain", "DELETE /api/users/{id}", "204", "Berhasil — data terhapus"],
            ["U-10", "Hapus user", "Menghapus akun sendiri", "DELETE /api/users/{id}", "422", "Gagal — tidak dapat menghapus akun sendiri"],
            ["U-11", "Kelola user", "Role bukan Admin", "GET /api/users", "403", "Gagal — tidak memiliki akses"],
            ["U-12", "Lihat daftar role", "Role Admin", "GET /api/roles", "200", "Berhasil — daftar role"],
            ["U-13", "Lihat daftar role", "Role bukan Admin", "GET /api/roles", "403", "Gagal — tidak memiliki akses"],
        ]),
    ]

    for title, rows in api_tables:
        add_heading_with_style(doc, title, 2)
        add_table_from_markdown(doc, rows, header=["ID", "Skenario", "Kondisi / Data", "Request", "Response", "Hasil"])

    # UAT Section
    add_heading_with_style(doc, "8. Skenario User Acceptance Test (UAT)", 1)
    doc.add_paragraph(
        "Skenario User Acceptance Test (UAT) berbasis skenario yang dijalankan oleh pengguna akhir untuk "
        "memvalidasi fungsionalitas sistem sesuai kebutuhan. Setiap skenario memuat langkah pengujian dan hasil "
        "yang diharapkan; kolom Hasil Aktual dan Status diisi oleh penguji (pengguna) selama uji penerimaan."
    )
    doc.add_paragraph("Skala Status: Lulus (sesuai harapan) / Gagal (tidak sesuai harapan).")

    add_uat_table(doc, "8.1 Login & Dashboard", [
        ["UAT-01", "Membuka halaman login", "Buka URL /login sebagai pengguna belum login", "Form login (email, password, ingat saya) ditampilkan", "", ""],
        ["UAT-02", "Login berhasil", "Isi email & password akun terdaftar, klik Login", "Redirect ke Dashboard /, sesi aktif", "", ""],
        ["UAT-03", "Login gagal (kredensial salah)", "Isi email/password salah", "Pesan 'Email atau password salah' ditampilkan, tetap di halaman login", "", ""],
        ["UAT-04", "Validasi form saat mengetik", "Kosongkan email / isi email tidak valid", "Error per field ditampilkan real-time, login diblokir", "", ""],
        ["UAT-05", "Indikator kekuatan password", "Ketik password lemah / kuat pada form login", "Meter kekuatan password & daftar aturan muncul real-time", "", ""],
        ["UAT-06", "Ingat saya", "Centang Ingat saya, login, tutup & buka browser", "Sesi tetap diingat (login otomatis)", "", ""],
        ["UAT-07", "Logout", "Klik ikon logout di header", "Sesi berakhir, redirect ke /login", "", ""],
        ["UAT-08", "Dashboard tampil", "Login dengan role apa pun, buka /", "Kartu statistik Disposisi & Kegiatan serta tabel Kegiatan per Periode tampil", "", ""],
        ["UAT-09", "Statistik disposisi akurat", "Bandingkan Total/Diterima/Ditolak/Diserahkan dengan data", "Jumlah sesuai data di halaman Disposisi", "", ""],
        ["UAT-10", "Statistik kegiatan akurat", "Bandingkan Total/Dilaksanakan/Tidak Dilaksanakan dengan data", "Jumlah sesuai data di halaman Kegiatan", "", ""],
        ["UAT-11", "Akses halaman tanpa login", "Buka / tanpa autentikasi", "Redirect ke /login", "", ""],
    ])

    add_uat_table(doc, "8.2 Kelola Surat (Staff)", [
        ["UAT-12", "Lihat daftar surat", "Login sebagai Staff, buka menu Surat", "Daftar surat ditampilkan", "", ""],
        ["UAT-13", "Tambah surat", "Isi form surat (valid), klik Simpan", "Surat tersimpan; otomatis dibuat Disposisi status Diterima & Pengingat untuk Asisten Daerah; redirect ke Disposisi", "", ""],
        ["UAT-14", "Tambah surat (field wajib kosong)", "Kosongkan field wajib", "Pesan error validasi, data tidak tersimpan", "", ""],
        ["UAT-15", "Ubah surat", "Ubah data surat, klik Simpan", "Data terupdate, toast berhasil", "", ""],
        ["UAT-16", "Hapus surat", "Klik Hapus, konfirmasi pada dialog", "Data terhapus, toast berhasil", "", ""],
        ["UAT-17", "Hak akses surat", "Login sebagai role selain Staff, buka menu Surat", "Menu Surat tidak tersedia", "", ""],
    ])

    add_uat_table(doc, "8.3 Kelola Disposisi (Staff / Asisten Daerah)", [
        ["UAT-18", "Lihat daftar disposisi", "Login sebagai Staff/Asisten, buka menu Disposisi", "Daftar disposisi (status & alasan) ditampilkan", "", ""],
        ["UAT-19", "Staff tidak dapat mengubah disposisi", "Staff membuka data disposisi", "Tombol Edit/aksi ubah tidak tersedia untuk Staff", "", ""],
        ["UAT-20", "Serahkan disposisi (Asisten)", "Asisten klik Menyerahkan", "Status menjadi Diserahkan", "", ""],
        ["UAT-20a", "Serahkan tanpa Penerima/Dituju", "Asisten klik Menyerahkan tanpa mengisi Penerima/Dituju", "Form menyerahkan (wajib Penerima & Dituju), simpan diblokir dengan error", "", ""],
        ["UAT-21", "Tolak disposisi tanpa alasan", "Asisten klik Menolak tanpa mengisi alasan", "Form menolak (wajib alasan), simpan diblokir dengan error", "", ""],
        ["UAT-22", "Tolak disposisi dengan alasan", "Asisten klik Menolak, isi alasan, simpan", "Status menjadi Ditolak + alasan tersimpan", "", ""],
        ["UAT-23", "Tidak ada aksi hapus disposisi", "Staff/Asisten mencari tombol Hapus pada disposisi", "Aksi hapus tidak tersedia untuk semua role", "", ""],
        ["UAT-24", "Asisten hanya dapat serahkan/tolak", "Asisten mencoba mengubah field data surat lain", "Hanya status serahkan/tolak (dan Penerima/Dituju/Alasan) yang dapat diubah", "", ""],
    ])

    add_uat_table(doc, "8.4 Kelola Kegiatan & Konfirmasi Kehadiran", [
        ["UAT-25", "Lihat daftar kegiatan", "Login sebagai Staff/Asisten/OPD, buka menu Kegiatan", "Daftar kegiatan + rekap hadir/tidak ditampilkan", "", ""],
        ["UAT-26", "Tambah kegiatan (jadwal kosong)", "Staff isi form kegiatan, klik Simpan", "Kegiatan tersimpan; otomatis dibuat Pengingat untuk semua role", "", ""],
        ["UAT-27", "Tambah kegiatan (jadwal bentrok)", "Staff isi kegiatan pada tanggal+jam yang sudah ada", "Kegiatan ditolak, pesan 'Jadwal bentrok' ditampilkan", "", ""],
        ["UAT-28", "Ubah kegiatan", "Staff ubah data kegiatan, simpan", "Data terupdate; jika bentrok, ditolak", "", ""],
        ["UAT-29", "Hapus kegiatan", "Staff klik Hapus, konfirmasi", "Data terhapus", "", ""],
        ["UAT-30", "Konfirmasi hadir (OPD)", "OPD klik Hadir pada kegiatan", "Kehadiran tercatat per akun OPD, dapat diubah", "", ""],
        ["UAT-31", "Konfirmasi tidak hadir (OPD)", "OPD klik Tidak Hadir", "Status kehadiran diperbarui", "", ""],
        ["UAT-32", "Lihat daftar OPD", "Staff/Asisten buka Daftar OPD pada kegiatan", "Daftar OPD yang mengonfirmasi hadir/tidak tampil", "", ""],
        ["UAT-33", "Role non-staff menambah kegiatan", "OPD/Asisten mencoba tambah kegiatan", "Tombol tambah/edit/hapus tidak tersedia", "", ""],
    ])

    add_uat_table(doc, "8.5 Kelola Pengingat (Staff / Asisten Daerah / OPD)", [
        ["UAT-34", "Lihat pengingat milik sendiri", "Login, buka menu Pengingat", "Hanya pengingat milik akun sendiri yang tampil", "", ""],
        ["UAT-35", "Tambah pengingat", "Isi judul, tanggal, prioritas, klik Simpan", "Pengingat tersimpan", "", ""],
        ["UAT-36", "Ubah pengingat", "Ubah data pengingat, simpan", "Data terupdate", "", ""],
        ["UAT-37", "Hapus pengingat", "Klik Hapus, konfirmasi", "Data terhapus", "", ""],
        ["UAT-38", "Pengingat milik user lain", "Buka/ubah/hapus pengingat milik user lain", "Diperlakukan sebagai tidak ditemukan (ditolak)", "", ""],
    ])

    add_uat_table(doc, "8.6 Kelola Pengguna (Admin)", [
        ["UAT-39", "Lihat daftar pengguna", "Login sebagai Admin, buka menu Pengguna", "Daftar pengguna + role ditampilkan", "", ""],
        ["UAT-40", "Tambah pengguna", "Isi data valid (password kuat), klik Simpan", "Pengguna tersimpan", "", ""],
        ["UAT-41", "Tambah pengguna (email duplikat)", "Isi email yang sudah terpakai", "Pesan error email duplikat, data tidak tersimpan", "", ""],
        ["UAT-42", "Tambah pengguna (password lemah)", "Isi password lemah", "Error validasi & indikator kekuatan password ditampilkan", "", ""],
        ["UAT-43", "Ubah pengguna", "Ubah nama/email/role/password, simpan", "Data terupdate", "", ""],
        ["UAT-44", "Hapus pengguna lain", "Hapus akun selain akun sendiri", "Data terhapus", "", ""],
        ["UAT-45", "Hapus akun sendiri", "Coba hapus akun yang sedang login", "Ditolak, pesan tidak dapat menghapus akun sendiri", "", ""],
        ["UAT-46", "Hak akses pengguna", "Login sebagai role selain Admin, buka menu Pengguna", "Menu Pengguna tidak tersedia", "", ""],
    ])

    # 6.7
    add_uat_table(doc, "8.7 Notifikasi Pengingat Real-Time (Staff / Asisten Daerah / OPD)", [
        ["UAT-47", "Lonceng notifikasi tampil", "Login sebagai Staff/Asisten/OPD", "Ikon lonceng Notifikasi Pengingat tampil di header (tidak tampil untuk Admin)", "", ""],
        ["UAT-48", "Notifikasi real-time dari Tambah Surat", "Staff menambah Surat; periksa akun Asisten Daerah pada tab lain (tanpa refresh)", "Lonceng bertambah real-time dengan badge jumlah belum dibaca", "", ""],
        ["UAT-49", "Notifikasi real-time dari Tambah Kegiatan", "Staff menambah Kegiatan; periksa akun lain pada tab lain (tanpa refresh)", "Lonceng bertambah real-time dengan badge jumlah belum dibaca", "", ""],
        ["UAT-50", "Pengingat manual tidak memicu notifikasi", "User menambah Pengingat manual pada halaman Pengingat", "Tidak ada lonceng/badge baru (sumber manual tidak dinotifikasikan)", "", ""],
        ["UAT-51", "Buka daftar notifikasi", "Klik lonceng notifikasi", "Dropdown berisi daftar notifikasi (label Surat/Kegiatan, waktu relatif, tanggal pengingat)", "", ""],
        ["UAT-52", "Baca satu notifikasi", "Klik salah satu notifikasi di dropdown", "Navigasi ke halaman Pengingat, badge unread berkurang satu", "", ""],
        ["UAT-53", "Tandai semua dibaca", "Klik Tandai semua dibaca pada dropdown", "Semua notifikasi berstatus dibaca, badge hilang, toast berhasil", "", ""],
        ["UAT-54", "Hak akses notifikasi", "Login sebagai Admin, lihat header", "Lonceng notifikasi tidak tersedia", "", ""],
    ])

    # 8. Blackbox Testing
    add_heading_with_style(doc, "9. Blackbox Testing", 1)
    doc.add_paragraph(
        "Pengujian blackbox dilakukan terhadap fungsi sistem tanpa memperhatikan struktur internal kode. "
        "Setiap skenario memetakan Kebutuhan Fungsional (FR) ke Skenario Uji, Langkah / Input, dan Hasil yang "
        "Diharapkan; kolom Hasil Aktual dan Status diisi oleh penguji."
    )
    doc.add_paragraph("Skala Status: Lulus (sesuai harapan) / Gagal (tidak sesuai harapan).")

    bb_rows = [
        ["BB-01", "FR-01", "Login berhasil", "Isi email & password valid, klik Login", "Redirect ke Dashboard /, sesi aktif", "", ""],
        ["BB-02", "FR-01", "Login gagal (kredensial salah)", "Isi email/password salah", "Pesan 'Email atau password salah', tetap di halaman login", "", ""],
        ["BB-03", "FR-01", "Login tanpa input", "Kosongkan field email/password", "Error validasi per field, login diblokir", "", ""],
        ["BB-04", "FR-01", "Logout", "Klik ikon logout di header", "Sesi berakhir, redirect ke /login", "", ""],
        ["BB-05", "FR-02", "Dashboard tampil", "Login dengan role apa pun, buka /", "Kartu statistik Disposisi & Kegiatan serta tabel Kegiatan per Periode tampil", "", ""],
        ["BB-06", "FR-03", "Tambah surat (valid)", "Staff isi form surat valid, klik Simpan", "Surat tersimpan; otomatis dibuat Disposisi (Diterima) & Pengingat ke Asisten Daerah", "", ""],
        ["BB-07", "FR-03", "Tambah surat (field wajib kosong)", "Staff kosongkan field wajib", "Pesan error validasi, data tidak tersimpan", "", ""],
        ["BB-08", "FR-03", "Ubah surat", "Staff ubah data surat, klik Simpan", "Data terupdate, toast berhasil", "", ""],
        ["BB-09", "FR-03", "Hapus surat", "Staff klik Hapus, konfirmasi pada dialog", "Data terhapus, toast berhasil", "", ""],
        ["BB-10", "FR-03", "Hak akses surat", "Login sebagai role selain Staff, buka menu Surat", "Menu Surat tidak tersedia", "", ""],
        ["BB-11", "FR-04", "Lihat daftar disposisi", "Login sebagai Staff/Asisten, buka menu Disposisi", "Daftar disposisi (status & alasan) ditampilkan", "", ""],
        ["BB-12", "FR-04", "Staff tidak dapat mengubah disposisi", "Staff membuka data disposisi", "Tombol Edit/aksi ubah tidak tersedia untuk Staff", "", ""],
        ["BB-13", "FR-05", "Serahkan disposisi", "Asisten klik Menyerahkan, isi Penerima & Dituju, simpan", "Status menjadi Diserahkan + Penerima/Dituju tersimpan + Pengingat otomatis ke Staff", "", ""],
        ["BB-14", "FR-05", "Serahkan tanpa Penerima/Dituju", "Asisten klik Menyerahkan tanpa mengisi Penerima/Dituju", "Form menyerahkan (wajib Penerima & Dituju), simpan diblokir", "", ""],
        ["BB-15", "FR-06", "Tolak disposisi dengan alasan", "Asisten klik Menolak, isi alasan, simpan", "Status menjadi Ditolak + alasan tersimpan + Pengingat otomatis ke Staff", "", ""],
        ["BB-16", "FR-06", "Tolak tanpa alasan", "Asisten klik Menolak tanpa mengisi alasan", "Form menolak (wajib alasan), simpan diblokir", "", ""],
        ["BB-17", "FR-06", "Tidak ada aksi hapus disposisi", "Staff/Asisten mencari tombol Hapus pada disposisi", "Aksi hapus tidak tersedia untuk semua role", "", ""],
        ["BB-18", "FR-07", "Tambah kegiatan (jadwal kosong)", "Staff isi form kegiatan valid, klik Simpan", "Kegiatan tersimpan; otomatis dibuat Pengingat ke semua role", "", ""],
        ["BB-19", "FR-07", "Tambah kegiatan (jadwal bentrok)", "Staff isi kegiatan pada tanggal+jam yang sama dengan kegiatan lain", "Kegiatan ditolak, pesan 'Jadwal bentrok' ditampilkan", "", ""],
        ["BB-20", "FR-07", "Ubah kegiatan", "Staff ubah data kegiatan, simpan", "Data terupdate; jika bentrok, ditolak", "", ""],
        ["BB-21", "FR-07", "Hapus kegiatan", "Staff klik Hapus, konfirmasi", "Data terhapus", "", ""],
        ["BB-22", "FR-07", "Hak akses kegiatan", "OPD/Asisten mencoba tambah kegiatan", "Tombol tambah/edit/hapus tidak tersedia", "", ""],
        ["BB-23", "FR-08", "Konfirmasi hadir", "OPD klik Hadir pada kegiatan", "Kehadiran tercatat per akun OPD, dapat diubah", "", ""],
        ["BB-24", "FR-08", "Konfirmasi tidak hadir", "OPD klik Tidak Hadir", "Status kehadiran diperbarui", "", ""],
        ["BB-25", "FR-08", "Ubah konfirmasi", "OPD konfirmasi ulang pada kegiatan yang sama", "Record kehadiran di-update", "", ""],
        ["BB-26", "FR-08", "Hak akses konfirmasi", "Role selain OPD mencoba konfirmasi kehadiran", "Aksi konfirmasi tidak tersedia/ditolak", "", ""],
        ["BB-27", "FR-09", "Lihat rekap & daftar OPD", "Staff/Asisten buka Daftar OPD pada kegiatan", "Rekap hadir/tidak serta daftar OPD yang mengonfirmasi tampil", "", ""],
        ["BB-28", "FR-10", "Tambah pengingat", "Isi judul, tanggal, prioritas, klik Simpan", "Pengingat tersimpan", "", ""],
        ["BB-29", "FR-10", "Ubah pengingat", "Ubah data pengingat, simpan", "Data terupdate", "", ""],
        ["BB-30", "FR-10", "Hapus pengingat", "Klik Hapus, konfirmasi", "Data terhapus", "", ""],
        ["BB-31", "FR-10", "Pengingat milik user lain", "Buka/ubah/hapus pengingat milik user lain", "Diperlakukan sebagai tidak ditemukan (ditolak)", "", ""],
        ["BB-32", "FR-11", "Notifikasi real-time dari Tambah Surat", "Staff menambah Surat; periksa akun Asisten Daerah pada tab lain (tanpa refresh)", "Lonceng bertambah real-time dengan badge jumlah belum dibaca", "", ""],
        ["BB-33", "FR-11", "Notifikasi real-time dari Tambah Kegiatan", "Staff menambah Kegiatan; periksa akun lain pada tab lain (tanpa refresh)", "Lonceng bertambah real-time dengan badge jumlah belum dibaca", "", ""],
        ["BB-34", "FR-11", "Notifikasi real-time saat Disposisi diserahkan/ditolak", "Asisten menyerahkan/menolak disposisi; periksa akun Staff pada tab lain (tanpa refresh)", "Lonceng bertambah real-time dengan badge jumlah belum dibaca", "", ""],
        ["BB-35", "FR-11", "Tandai notifikasi dibaca", "Klik salah satu notifikasi di dropdown", "Navigasi ke halaman Pengingat, badge unread berkurang satu", "", ""],
        ["BB-36", "FR-11", "Tandai semua dibaca", "Klik Tandai semua dibaca pada dropdown", "Semua notifikasi berstatus dibaca, badge hilang", "", ""],
        ["BB-37", "FR-11", "Hak akses notifikasi", "Login sebagai Admin, lihat header", "Lonceng notifikasi tidak tersedia", "", ""],
        ["BB-38", "FR-12", "Tambah pengguna (valid)", "Admin isi data valid (password kuat), klik Simpan", "Pengguna tersimpan", "", ""],
        ["BB-39", "FR-12", "Tambah pengguna (email duplikat)", "Isi email yang sudah terpakai", "Pesan error email duplikat, data tidak tersimpan", "", ""],
        ["BB-40", "FR-12", "Tambah pengguna (password lemah)", "Isi password lemah", "Error validasi & indikator kekuatan password ditampilkan", "", ""],
        ["BB-41", "FR-12", "Hapus akun sendiri", "Coba hapus akun yang sedang login", "Ditolak, pesan tidak dapat menghapus akun sendiri", "", ""],
        ["BB-42", "FR-12", "Hak akses pengguna", "Login sebagai role selain Admin, buka menu Pengguna", "Menu Pengguna tidak tersedia", "", ""],
    ]
    add_table_from_markdown(doc, bb_rows, header=["ID", "Kebutuhan Fungsional", "Skenario Uji", "Langkah / Input", "Hasil yang Diharapkan", "Hasil Aktual", "Status"])

    # Save the standalone UML documentation document.
    docx_path = Path("c:/Users/V14/SI_Pengelolaan_Data_Agenda/laravel-vue-mvc/docs/UML_Documentation.docx")
    doc.save(str(docx_path))
    print(f"Successfully generated {docx_path}")
    print("  - Standalone UML documentation (Login/Dashboard diagrams + API scenarios + UAT + Blackbox Testing)")


if __name__ == "__main__":
    main()